import xlrd
import xlwt
import argparse
import re
import pyperclip
import os
from docx import Document

#-i "C:\Users\kanka\Downloads\Python_MCQ.xls" -o "ExitFile.xls" -wsn 3 -ln 25

###################################################################################Argument parser can be commented see below########################
input_Filename=''
output_Filename=''
ws_Number=0
Linenumbers=1000
copyflag=0
helptext="-i <Input Excel FilePath> \n -o <Input Excel FilePath> \n -wsn <Worksheet Number> \n -ln(Optional) <How many Rows to parse>"

parser=argparse.ArgumentParser(description="Converts METL Code to EDUOnline Format")
parser.add_argument('-help',dest='',type=str,nargs='+',help=helptext)
parser.add_argument('-i',dest='input',type=str,nargs='+',help='-i <Input Excel FilePath>')
parser.add_argument('-o',dest='output',type=str,nargs='+',help='-o <Input Excel FilePath>')
parser.add_argument('-c',dest='copy',type=int,nargs='+',help='-c 1 or 0 , 1 --> Enable clipboard copy(default) , 0 Clipboard copy disabled')
parser.add_argument('-wsn',dest='wsn',type=int,nargs='+',help='-wsn(Default is Sheet 1) <Worksheet Number>')
parser.add_argument('-ln',dest='linecount',type=int,nargs='+',help='-ln(Optional, default is 1000) <How many Rows to parse>')

args=parser.parse_args()

try:
    input_Filename=args.input[0]
except Exception:
    print("Opps! Did you forget to put The Input File Argument?")
try:
    output_Filename=args.output[0]
    ws_Number=args.wsn[0]
    Linenumbers=args.linecount[0]
    copyflag=args.copy[0]
except Exception:
    pass

if input_Filename!='':
    output_Filename=os.path.basename(input_Filename).split(".")[0]+".docx"


def ExcelReader(filename,WorksheetName="",type=0): #Default type=0 MCQ , Type2 CheckBox , Type3 Short Answer
    workbook = xlrd.open_workbook(filename)
    if WorksheetName:
        worksheet = workbook.sheet_by_name(WorksheetName)
    else:
        print("Excel Sheet is Empty")
        return [0,0,0]

    AllData=[]
    Difficulty=[]
    Questions = []
    Option1 = []
    Option2 = []
    Option3 = []
    Option4 = []
    Level1_Data = []
    Level2_Data = []
    Level3_Data = []
    DocumentSchema={}
    headers = [str(cell.value) for cell in worksheet.row(0)]
    for columns in headers:
        if columns.lower().find("difficulty")!=-1:
            DocumentSchema["difficulty"]=headers.index(columns)
        if columns.lower().find("question text")!=-1:
            DocumentSchema["question"]=headers.index(columns)
        if columns.lower().find("1") != -1:
            DocumentSchema["ch1"] = headers.index(columns)
        if columns.lower().find("2") !=-1:
            DocumentSchema["ch2"] = headers.index(columns)
        if columns.lower().find("3") !=-1:
            DocumentSchema["ch3"] = headers.index(columns)
        if columns.lower().find("4") !=-1:
            DocumentSchema["ch4"] = headers.index(columns)
        if columns.lower().find("5") !=-1:
            DocumentSchema["ch5"] = headers.index(columns)
        if columns.lower().find("correct") !=-1:
            DocumentSchema["answer"] = headers.index(columns)
    print(DocumentSchema)
    try:
        for i in range(Linenumbers):
            AllData.append(worksheet.row_values(i))
    except Exception:
        pass
    for i in AllData:
        if AllData.index(i) != 0:
            _temp=[]
            _temp.append(str(i[DocumentSchema['question']]))
            if type<2:
                for counter in range(1,6):
                    # print(i[DocumentSchema["ch" + str(counter)]])
                    if i[DocumentSchema['answer']].find(str(counter))!=-1:
                        if type==0:
                            _temp.append("(x) "+str(i[DocumentSchema["ch"+str(counter)]]))
                        elif type==1:
                            _temp.append("[x] " + str(i[DocumentSchema["ch" + str(counter)]]))

                    elif i[DocumentSchema["ch"+str(counter)]]!='':
                        if type == 0:
                            _temp.append("( ) " + str(i[DocumentSchema["ch" + str(counter)]]))
                        elif type == 1:
                            _temp.append("[ ] " + str(i[DocumentSchema["ch" + str(counter)]]))
                # print(_temp)
                if i[DocumentSchema["difficulty"]].lower().find("simple")!=-1 or i[DocumentSchema["difficulty"]].lower().find("easy")!=-1:
                    # print("Easy")
                    Level1_Data.append(_temp)
                if i[DocumentSchema["difficulty"]].lower().find("medium")!=-1:
                    # print("medium")
                    Level2_Data.append(_temp)
                if i[DocumentSchema["difficulty"]].lower().find("hard")!=-1 or i[DocumentSchema["difficulty"]].lower().find("difficult")!=-1:
                    Level3_Data.append(_temp)
                    # print("hard")
            else:
                index_question=DocumentSchema['question']
                print(i[DocumentSchema["difficulty"]].lower())
                for counter in range(1,6):
                    if counter==1 and i[index_question+counter]!='' :
                        _temp.append("= "+i[index_question+counter])
                    else:
                        if i[index_question+counter]!='':
                            _temp.append("or= " + i[index_question+counter])
                if i[DocumentSchema["difficulty"]].lower().find("simple") != -1 or i[
                    DocumentSchema["difficulty"]].lower().find("easy") != -1:
                    # print("Easy")
                    Level1_Data.append(_temp)
                if i[DocumentSchema["difficulty"]].lower().find("medium") != -1:
                    # print("medium")
                    Level2_Data.append(_temp)
                if i[DocumentSchema["difficulty"]].lower().find("hard") != -1 or i[
                    DocumentSchema["difficulty"]].lower().find("difficult") != -1:
                    Level3_Data.append(_temp)
    return [Level1_Data,Level2_Data,Level3_Data]



document = Document()
book = xlrd.open_workbook(input_Filename)
for sheet in book.sheets():
    if sheet.name.lower().find("mcq")==0:
        print("Processing Sheet",sheet.name)
        a1, b1, c1 = ExcelReader(input_Filename, sheet.name, 0)
        if a1 !=[]:
            document.add_heading('MCQ', 0)
            document.add_heading('MCQ Simple', 2)
            for i in a1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_paragraph("\n")
            document.add_page_break()
        if b1 !=[]:
            document.add_heading('MCQ Medium', 2)
            for i in b1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if c1 !=[]:
            document.add_heading('MCQ Difficult', 2)
            for i in c1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
    elif sheet.name.lower().find("check")==0:
        print("Processing Sheet",sheet.name)
        a1, b1, c1 = ExcelReader(input_Filename, sheet.name, 1)
        if a1 !=[]:
            document.add_heading('Checkbox', 0)
            document.add_heading('Checkbox Simple', 2)
            for i in a1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if b1 !=[]:
            document.add_heading('Checkbox Medium', 2)
            for i in b1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if c1 !=[]:
            document.add_heading('Checkbox Difficult', 2)
            for i in c1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
    elif sheet.name.lower().find("short")==0:
        print("Processing Sheet",sheet.name)
        a1, b1, c1= ExcelReader(input_Filename, sheet.name, 2)
        if a1 !=[]:
            document.add_heading('ShortAnswers', 0)
            document.add_heading('ShortAnswers Simple', 2)
            for i in a1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if b1 !=[]:
            document.add_heading('ShortAnswers Medium', 2)
            for i in b1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if c1 !=[]:
            document.add_heading('ShortAnswers Difficult', 2)
            for i in c1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()

    else:
        print("Invalid Sheet ",sheet.name)
document.save(output_Filename)