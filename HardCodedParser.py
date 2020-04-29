import xlrd
import xlwt
import argparse
import re
import pyperclip
import os
from docx import Document

#-i "C:\Users\kanka\Downloads\Python_MCQ.xls" -o "ExitFile.xls" -wsn 3 -ln 25

###################################################################################Argument parser can be commented see below########################
input_Filename=''
output_Filename=''
ws_Number=0
Linenumbers=1000
copyflag=0
helptext="-i <Input Excel FilePath> \n -o <Input Excel FilePath> \n -wsn <Worksheet Number> \n -ln(Optional) <How many Rows to parse>"

parser=argparse.ArgumentParser(description="Converts METL Code to EDUOnline Format")
parser.add_argument('-help',dest='',type=str,nargs='+',help=helptext)
parser.add_argument('-i',dest='input',type=str,nargs='+',help='-i <Input Excel FilePath>')
parser.add_argument('-o',dest='output',type=str,nargs='+',help='-o <Input Excel FilePath>')
parser.add_argument('-c',dest='copy',type=int,nargs='+',help='-c 1 or 0 , 1 --> Enable clipboard copy(default) , 0 Clipboard copy disabled')
parser.add_argument('-wsn',dest='wsn',type=int,nargs='+',help='-wsn(Default is Sheet 1) <Worksheet Number>')
parser.add_argument('-ln',dest='linecount',type=int,nargs='+',help='-ln(Optional, default is 1000) <How many Rows to parse>')

args=parser.parse_args()

try:
    input_Filename=args.input[0]
except Exception:
    print("Opps! Did you forget to put The Input File Argument?")
try:
    output_Filename=args.output[0]
    ws_Number=args.wsn[0]
    Linenumbers=args.linecount[0]
    copyflag=args.copy[0]
except Exception:
    pass

if input_Filename!='':
    output_Filename=os.path.basename(input_Filename).split(".")[0]+".docx"
    print(output_Filename)


def ExcelReader(filename,WorksheetName="",type=0): #Default type=0 MCQ , Type2 CheckBox , Type3 Short Answer
    workbook = xlrd.open_workbook(filename)
    if WorksheetName:
        worksheet = workbook.sheet_by_name(WorksheetName)
    else:
        return [0,0,0]
        print("Excel Sheet is Empty")
    AllData=[]
    Difficulty=[]
    Questions = []
    Option1 = []
    Option2 = []
    Option3 = []
    Option4 = []
    Level1_Data = []
    Level2_Data = []
    Level3_Data = []
    try:
        for i in range(Linenumbers):
            AllData.append(worksheet.row_values(i))
    except Exception:
        pass
    for i in AllData:
        print(i)
        if AllData.index(i) != 0:
            Questions.append(i[2])
            Difficulty.append(i[1].lower())
            if type>=2:
                _temp = []
                if i!='':
                    _temp.append(i[2])
                if str(i[3])!='':
                    Option1.append("= "+str(i[3]))
                    _temp.append("= " + str(i[3]))
                if str(i[4])!='':
                    Option2.append("or= "+str(i[4]))
                    _temp.append("or= " + str(i[4]))
                if str(i[5])!='':
                    Option3.append("or= "+str(i[5]))
                    _temp.append("or= " + str(i[5]))
                if str(i[6])!='':
                    Option4.append("or= "+str(i[6]))
                    _temp.append("or= " + str(i[6]))

            else:
                if i[8].find("1") != -1:
                    # print('Option 1 is correct')
                    if type==0:
                        Option1.append("(x) " + str(i[3]))
                    elif type==1:
                        Option1.append("[x] " + str(i[3]))
                else:
                    if type == 0:
                        Option1.append("( ) " + str(i[3]))
                    elif type == 1:
                        Option1.append("[ ] " + str(i[3]))

                if i[8].find("2") != -1:
                    # print('Option 1 is correct')
                    if type == 0:
                        Option2.append("(x) " + str(i[4]))
                    elif type == 1:
                        Option2.append("[x] " + str(i[4]))
                else:
                    if type == 0:
                        Option2.append("( ) " + str(i[4]))
                    elif type == 1:
                        Option2.append("[ ] " + str(i[4]))

                if i[8].find("3") != -1:
                    # print('Option 1 is correct')
                    if type == 0:
                        Option3.append("(x) " + str(i[5]))
                    elif type == 1:
                        Option3.append("[x] " + str(i[5]))
                else:
                    if type == 0:
                        Option3.append("( ) " + str(i[5]))
                    elif type == 1:
                        Option3.append("[ ] " + str(i[5]))

                if i[8].find("4") != -1:
                    # print('Option 1 is correct')
                    if type == 0:
                        Option4.append("(x) " + str(i[6]))
                    elif type == 1:
                        Option4.append("[x] " + str(i[6]))
                else:
                    if type == 0:
                        Option4.append("( ) " + str(i[6]))
                    elif type == 1:
                        Option4.append("[ ] " + str(i[6]))
    if type>=2:

        for j in Questions:
            if j != '':
                if Difficulty[Questions.index(j)].find("simple") == 0 or Difficulty[Questions.index(j)].find("easy") == 0:
                    Level1_Data.append(_temp)
                if Difficulty[Questions.index(j)].find("medium") == 0:
                    Level2_Data.append(_temp)
                if Difficulty[Questions.index(j)].find("hard") == 0 or Difficulty[Questions.index(j)].find(
                        "difficult") == 0:
                    Level3_Data.append(_temp)
    if type<2:
        for j in Questions:
            if j != '':
                 if Difficulty[Questions.index(j)].find("simple") == 0 or Difficulty[Questions.index(j)].find(
                        "easy") == 0:
                    Level1_Data.append([j, Option1[Questions.index(j)], Option2[Questions.index(j)],
                                        Option3[Questions.index(j)], Option4[Questions.index(j)]])
                 elif Difficulty[Questions.index(j)].find("medium") == 0:
                    Level2_Data.append([j, Option1[Questions.index(j)], Option2[Questions.index(j)],
                                        Option3[Questions.index(j)], Option4[Questions.index(j)]])
                 elif Difficulty[Questions.index(j)].find("hard") == 0 or Difficulty[Questions.index(j)].find(
                        "difficult") == 0:
                    Level3_Data.append([j, Option1[Questions.index(j)], Option2[Questions.index(j)],
                                        Option3[Questions.index(j)], Option4[Questions.index(j)]])
    return [Level1_Data,Level2_Data,Level3_Data]


document = Document()
book = xlrd.open_workbook(input_Filename)
for sheet in book.sheets():
    if sheet.name.lower().find("mcq")==0:
        print("Processing Sheet",sheet.name)
        a1, b1, c1 = ExcelReader(input_Filename, sheet.name, 0)
        if a1 !=[]:
            document.add_heading('MCQ', 0)
            document.add_heading('MCQ Simple', 2)
            for i in a1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if b1 !=[]:
            document.add_heading('MCQ Medium', 2)
            for i in b1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if c1 !=[]:
            document.add_heading('MCQ Difficult', 2)
            for i in c1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
    elif sheet.name.lower().find("check")==0:
        print("Processing Sheet",sheet.name)
        a1, b1, c1 = ExcelReader(input_Filename, sheet.name, 1)
        if a1 !=[]:
            document.add_heading('Checkbox', 0)
            document.add_heading('Checkbox Simple', 2)
            for i in a1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if b1 !=[]:
            document.add_heading('Checkbox Medium', 2)
            for i in b1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if c1 !=[]:
            document.add_heading('Checkbox Difficult', 2)
            for i in c1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
    elif sheet.name.lower().find("short")==0:
        print("Processing Sheet",sheet.name)
        a1, b1, c1= ExcelReader(input_Filename, sheet.name, 2)
        if a1 !=[]:
            document.add_heading('ShortAnswers', 0)
            document.add_heading('ShortAnswers Simple', 2)
            for i in a1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if b1 !=[]:
            document.add_heading('ShortAnswers Medium', 2)
            for i in b1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()
        if c1 !=[]:
            document.add_heading('ShortAnswers Difficult', 2)
            for i in c1:
                for lines in i:
                    document.add_paragraph(lines)
            document.add_page_break()

    else:
        print("Invalid Sheet ",sheet.name)
document.save(output_Filename)